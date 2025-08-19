# 실행
# python extractor.py

import os
import tempfile
import shutil
import time
from typing import Literal, List
import gc
import re

import fitz  # PyMuPDF
from PIL import Image  # noqa: F401 (미래 전처리용)
from docx import Document
import pytesseract

from backend.extract.text_parser import extract_text_from_docx, extract_text_from_pdf
from backend.extract.ocr_parser import extract_text_from_image
from backend.extract.utils import get_file_type
from backend.extract.vlm_parser import vlm_extract_caption_bytes  # ⬅️ VLM 캡션 함수 사용

SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9가-힣._-]+")

# -------------------- 환경/설정 --------------------
# Windows Tesseract 경로 (환경에 맞게 조정)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# config에서 가져오되, 없으면 안전 기본값 사용
try:
    from config import PDF_DPI, OCR_MAX_PAGES, OCR_AUTO_SKIP_TEXT_LEN, OCR_THREAD_WORKERS
except Exception:
    PDF_DPI = 120                # 300→120로 낮춰 메모리/디스크 부담 감소
    OCR_MAX_PAGES = 10           # OCR 최대 페이지 수 제한
    OCR_AUTO_SKIP_TEXT_LEN = 800 # 텍스트가 충분하면 OCR 생략
    OCR_THREAD_WORKERS = 2       # 병렬 OCR 스레드 (Windows는 2~3 권장)

# (선택) OpenMP/Tesseract 내부 스레드 제한 – 과도한 병렬 방지
os.environ.setdefault("OMP_THREAD_LIMIT", "1")


# -------------------- 공용 유틸 --------------------
def _cleanup_paths(paths: List[str]):
    """임시파일/폴더 정리. 우리가 만든 tmpdir 패턴이면 폴더ごと 삭제."""
    for p in set(paths):
        try:
            if not os.path.exists(p):
                continue
            base = os.path.dirname(p)
            # 아래 prefix는 본 파일에서 생성하는 임시폴더 접두사
            if os.path.basename(base).startswith(("pdfimg_", "docximg_")):
                shutil.rmtree(base, ignore_errors=True)
            else:
                os.remove(p)
        except Exception as e:
            print("[CLEANUP] error:", e)


# -------------------- DOCX 이미지 추출 --------------------
def _extract_images_from_docx(path: str) -> List[str]:
    """
    DOCX 내 이미지를 임시파일로 추출(경로 리스트 반환).
    메모리에 이미지를 오래 들고 있지 않고 파일로만 다룸.
    """
    tmpdir = tempfile.mkdtemp(prefix="docximg_")
    out: List[str] = []
    try:
        doc = Document(path)
        rels = list(doc.part._rels.values())
        for idx, rel in enumerate(rels):
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                img_path = os.path.join(tmpdir, f"img_{idx}.png")
                with open(img_path, "wb") as f:
                    f.write(img_data)
                out.append(img_path)
    except Exception as e:
        print("[DOCX IMG] error:", e)
    return out


# -------------------- PDF → 이미지 (PyMuPDF) --------------------
def _pdf_to_image_paths(pdf_path: str, max_pages: int, dpi: int) -> List[str]:
    """
    PyMuPDF로 PDF 각 페이지를 PNG 파일로 렌더링하고 경로 리스트를 반환.
    *서브프로세스 사용 안 함* → stdout/stderr 버퍼 폭주 원천 차단.
    """
    if max_pages <= 0:
        return []
    tmpdir = tempfile.mkdtemp(prefix="pdfimg_")
    paths: List[str] = []
    try:
        doc = fitz.open(pdf_path)
        total = min(len(doc), max_pages)
        # 72dpi 기준 → 스케일 행렬
        zoom = float(dpi) / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for i in range(total):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=mat, alpha=False)  # RGB
            out_path = os.path.join(tmpdir, f"page_{i+1}.png")
            pix.save(out_path)
            paths.append(out_path)
        doc.close()
    except Exception as e:
        print("[PDF->IMG(PyMuPDF)] error:", e)
    return paths


# -------------------- PDF 내장(embedded) 이미지만 개별 추출 --------------------
def _extract_images_from_pdf_embedded(pdf_path: str, max_pages: int = 0) -> List[str]:
    """
    PDF 페이지에 '내장된 이미지'들을 각각 파일로 추출해 경로 리스트를 반환.
    - 페이지 전체 스냅샷이 아니라, 이미지 오브젝트(xref) 단위 추출
    - 추출 즉시 디스크에 저장하므로 메모리 점유 최소화
    """
    tmpdir = tempfile.mkdtemp(prefix="pdfimg_")
    out: List[str] = []
    try:
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        last = page_count if max_pages <= 0 else min(page_count, max_pages)
        seen_xref = set()
        for i in range(last):
            page = doc.load_page(i)
            # full=True로 해야 smask 등 정보 포함
            for img_info in page.get_images(full=True):
                try:
                    xref = img_info[0]
                    if xref in seen_xref:
                        continue
                    seen_xref.add(xref)
                    # xref로 이미지 추출
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n >= 5:  # CMYK/GRAY+ALPHA 등 → RGB 변환
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    out_path = os.path.join(tmpdir, f"p{i+1}_xref{xref}.png")
                    pix.save(out_path)
                    out.append(out_path)
                    # 메모리 해제
                    pix = None
                except Exception as e:
                    print(f"[PDF EMBED IMG] extract error on page {i+1}: {e}")
                    continue
        doc.close()
    except Exception as e:
        print("[PDF EMBED IMG] open error:", e)
    return out
    


# ========== 새 핵심 함수: 텍스트+VLM을 .txt로 직저장 (무파일 이미지 처리) ==========

def _iter_all_images(filepath: str, max_pages: int, dpi: int, prefer_embedded: bool = True, fallback_render: bool = True):
    """
    파일 내 모든 이미지 경로를 '생성기'로 반환.
    - PDF:
        prefer_embedded=True  → 내장 이미지(xref)만 개별 추출 (메모리 적게)
        fallback_render=True  → 내장 이미지가 1장도 없으면 페이지 렌더로 대체
    - DOCX: 임베디드 이미지만 추출
    """
    ftype = get_file_type(filepath)
    if ftype == "pdf":
        if prefer_embedded:
            embedded = _extract_images_from_pdf_embedded(filepath, max_pages=max_pages)
            if embedded:
                for p in embedded:
                    yield p
                return
            # 내장 이미지가 없다면 필요 시 페이지 렌더 fallback
            if not embedded and fallback_render:
                for p in _pdf_to_image_paths(filepath, max_pages=max_pages, dpi=dpi):
                    yield p
        else:
            for p in _pdf_to_image_paths(filepath, max_pages=max_pages, dpi=dpi):
                yield p
    elif ftype == "docx":
        for p in _extract_images_from_docx(filepath):
            yield p
    else:
        return

# -------------------- 핵심: 추출 파이프라인 --------------------
def extract_all(
    filepath: str,
    use_ocr: Literal["auto", True, False] = "auto",
    use_vlm: bool = False,  # VLM은 매우 무거움. 기본 False.
) -> dict:
    """
    텍스트 우선 → 필요할 때만 OCR.
    이미지/임시파일은 경로만 다루고, 처리 후 즉시 정리.
    """
    ftype = get_file_type(filepath)
    result = {"text": "", "ocr": [], "vlm": []}

    # 1) 텍스트 우선 추출
    t0 = time.time()
    if ftype == "docx":
        result["text"] = extract_text_from_docx(filepath)
        doc_images = _extract_images_from_docx(filepath)
        pdf_images = []
    elif ftype == "pdf":
        result["text"] = extract_text_from_pdf(filepath)
        doc_images = []
        pdf_images = _pdf_to_image_paths(filepath, max_pages=OCR_MAX_PAGES, dpi=PDF_DPI)
    else:
        raise ValueError("지원하지 않는 파일 형식입니다.")
    print(f"[TIME] 텍스트 추출만 소요: {time.time() - t0:.2f}s, text_len={len(result['text'])}")

    # 2) OCR 필요 여부 판단
    need_ocr = (use_ocr is True) or (use_ocr == "auto" and len(result["text"]) < OCR_AUTO_SKIP_TEXT_LEN)

    if need_ocr:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        targets: List[str] = []
        targets.extend(pdf_images)
        targets.extend(doc_images)

        t1 = time.time()
        if targets:
            ocr_results: List[str] = []
            # 병렬도는 낮게 (윈도우 안정성)
            with ThreadPoolExecutor(max_workers=OCR_THREAD_WORKERS) as ex:
                futs = [ex.submit(extract_text_from_image, p) for p in targets]
                for fu in as_completed(futs):
                    try:
                        ocr_results.append(fu.result())
                    except Exception as e:
                        print("[OCR] error:", e)
                        ocr_results.append("")
            result["ocr"] = ocr_results
        print(f"[TIME] OCR 소요: {time.time() - t1:.2f}s (pages={len(targets)})")

    # 3) VLM (OFF 권장) — 필요 시 여기에 추가
    if use_vlm:
        pass

    # 4) 임시파일 정리
    _cleanup_paths(pdf_images + doc_images)

    return result





SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9가-힣._-]+")

def _slugify(s: str, max_len: int = 80) -> str:
    s = (s or "").strip().replace(" ", "_")
    s = SAFE_NAME_RE.sub("_", s)
    return s[:max_len] if s else "untitled"

# ========== 내장 이미지 바이트 제너레이터 ==========
def _iter_pdf_embedded_image_bytes(pdf_path: str, max_pages: int = 0):
    """
    PDF 내장 이미지(xref)를 '한 장씩 bytes'로 yield.
    디스크에 파일 저장하지 않음. 처리 후 즉시 해제 가능.
    """
    try:
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        last = page_count if max_pages <= 0 else min(page_count, max_pages)
        seen_xref = set()
        for i in range(last):
            page = doc.load_page(i)
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                if xref in seen_xref:
                    continue
                seen_xref.add(xref)
                try:
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n >= 5:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    # PNG 바이트
                    img_bytes = pix.tobytes("png")
                    yield (i + 1, xref, img_bytes)
                except Exception as e:
                    print(f"[PDF EMBED IMG] page {i+1} xref {xref} error:", e)
                finally:
                    # Pixmap 메모리 해제
                    try:
                        del pix
                    except:
                        pass
                    gc.collect()
        doc.close()
    except Exception as e:
        print("[PDF EMBED IMG] open error:", e)

def _iter_docx_embedded_image_bytes(path: str):
    """
    DOCX 내장 이미지를 '한 장씩 bytes'로 yield.
    파일 저장 없이 바로 VLM에 전달 가능.
    """
    try:
        doc = Document(path)
        rels = list(doc.part._rels.values())
        for idx, rel in enumerate(rels, start=1):
            if "image" in rel.target_ref:
                try:
                    img_bytes = rel.target_part.blob
                    yield (idx, img_bytes)
                except Exception as e:
                    print(f"[DOCX IMG] error idx={idx}:", e)
                finally:
                    gc.collect()
    except Exception as e:
        print("[DOCX IMG] open error:", e)

# ========== 텍스트 스트리밍 기록 ==========
def _write_text_stream_pdf(pdf_path: str, wf, max_pages: int | None = None):
    """
    PDF 텍스트를 페이지별로 추출해 파일에 바로 기록 (대용량 안전).
    extract_text_from_pdf가 한방에 메모리에 올리면 메모리 부담이 커서,
    여기선 페이지 단위로 처리.
    """
    try:
        doc = fitz.open(pdf_path)
        last = len(doc) if not max_pages or max_pages <= 0 else min(len(doc), max_pages)
        for i in range(last):
            page = doc.load_page(i)
            txt = page.get_text() or ""
            if txt:
                wf.write(txt)
                wf.write("\n")
            # 즉시 해제
            del txt, page
            gc.collect()
        doc.close()
    except Exception as e:
        # fallback: 기존 함수 (어쩔 수 없을 때만)
        txt = extract_text_from_pdf(pdf_path) or ""
        wf.write(txt + "\n")
        del txt
        gc.collect()

def _write_text_stream_docx(path: str, wf):
    """
    DOCX 텍스트를 문단 단위로 스트리밍 기록.
    """
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            t = p.text or ""
            if t:
                wf.write(t)
                wf.write("\n")
            del t
            gc.collect()
    except Exception:
        txt = extract_text_from_docx(path) or ""
        wf.write(txt + "\n")
        del txt
        gc.collect()

# ========== 새 핵심 함수: 텍스트+VLM을 .txt로 직저장 (무파일 이미지 처리) ==========
def extract_to_txt(
    filepath: str,
    out_txt_path: str | None = None,
    use_ocr: bool = False,      # 요청대로 기본은 VLM만
    use_vlm: bool = True,
    vlm_prompt: str = "이미지의 핵심 내용/구조/수치/함의를 간결히 요약해 주세요. 표/그래프라면 축/단위/범례도 함께.",
    max_pages: int | None = None,
    doc_type_for_name: str | None = None,  # "RFP" / "Proposal" 등 파일명에 반영
    title_for_name: str | None = None,     # 사람이 읽을 제목 반영
    doc_id_for_name: str | None = None,    # UUID 등
) -> dict:
    """
    - 텍스트: 스트리밍으로 바로 .txt 기록 후 즉시 메모리 해제
    - 이미지: 내장 이미지 바이트를 1장씩 VLM에 전달 → 캡션을 txt에 append
        * 이미지 파일로 저장하지 않음
    - PDF: 내장 이미지가 없으면 '페이지 렌더링 fallback' 없이 이미지 없음 처리
    - out: {"txt_path", "text_len", "n_images", "n_vlm_ok"}
    """
    t0 = time.time()
    ftype = get_file_type(filepath)

    # 출력 파일명 자동 구성 (요구사항: docType이 드러나도록)
    if out_txt_path is None:
        base_dir = os.path.dirname(filepath)
        parts = []
        if doc_type_for_name:
            parts.append(_slugify(doc_type_for_name))
        if title_for_name:
            parts.append(_slugify(title_for_name))
        if doc_id_for_name:
            parts.append(_slugify(doc_id_for_name))
        if not parts:
            parts = ["document"]
        out_txt_name = "__".join(parts) + ".extracted.txt"
        out_txt_path = os.path.join(base_dir, out_txt_name)

    text_len = 0
    n_images = 0
    n_vlm_ok = 0

    with open(out_txt_path, "w", encoding="utf-8") as wf:
        # 1) 텍스트 본문 (스트리밍)
        wf.write("[[TEXT_BEGIN]]\n")
        if ftype == "pdf":
            _write_text_stream_pdf(filepath, wf, max_pages=max_pages or OCR_MAX_PAGES)
        elif ftype == "docx":
            _write_text_stream_docx(filepath, wf)
        else:
            raise ValueError("지원하지 않는 파일 형식입니다.")
        wf.write("[[TEXT_END]]\n")

        # 파일 길이(텍스트 길이) 계산: 파일 포인터로 측정
        wf.flush()
        try:
            text_len = os.path.getsize(out_txt_path)
        except Exception:
            text_len = 0

        # 2) 이미지 → VLM (내장 이미지만, fallback 렌더링 없음)
        if use_vlm:
            if ftype == "pdf":
                img_iter = _iter_pdf_embedded_image_bytes(filepath, max_pages=max_pages or OCR_MAX_PAGES)
                for page_idx, xref, img_bytes in img_iter:
                    n_images += 1
                    caption = ""
                    try:
                        caption = vlm_extract_caption_bytes(img_bytes, vlm_prompt)
                        n_vlm_ok += 1 if caption else 0
                    except Exception as e:
                        print(f"[VLM] error p{page_idx}/xref{xref}:", e)
                    # 기록
                    wf.write(f"\n[[IMAGE_{n_images}_BEGIN p{page_idx} xref{xref}]]\n")
                    wf.write(caption.strip() if caption else "(no-caption)")
                    wf.write(f"\n[[IMAGE_{n_images}_END]]\n")
                    # 즉시 메모리 해제
                    del img_bytes, caption
                    gc.collect()

            elif ftype == "docx":
                img_iter = _iter_docx_embedded_image_bytes(filepath)
                for idx, img_bytes in img_iter:
                    n_images += 1
                    caption = ""
                    try:
                        caption = vlm_extract_caption_bytes(img_bytes, vlm_prompt)
                        n_vlm_ok += 1 if caption else 0
                    except Exception as e:
                        print(f"[VLM] error docx img {idx}:", e)
                    wf.write(f"\n[[IMAGE_{n_images}_BEGIN docxImg{idx}]]\n")
                    wf.write(caption.strip() if caption else "(no-caption)")
                    wf.write(f"\n[[IMAGE_{n_images}_END]]\n")
                    del img_bytes, caption
                    gc.collect()

    print(f"[TIME] extract_to_txt: {time.time()-t0:.2f}s | text_bytes~={text_len} | images={n_images} | vlm_ok={n_vlm_ok} | out={out_txt_path}")
    return {"txt_path": out_txt_path, "text_len": text_len, "n_images": n_images, "n_vlm_ok": n_vlm_ok}


# -------------------- 단독 실행 테스트 --------------------
if __name__ == "__main__":
    # 환경에 맞게 경로 조정
    path = "C:/Users/KJY/rfp-analyzer/backend/extract/test_files/sample.docx"
    print("Trying to load file:", path)
    res = extract_all(path, use_ocr="auto", use_vlm=False)
    print("\n[텍스트](앞부분):\n", res["text"][:500])
    print("\n[OCR 샘플]:\n", res["ocr"][0] if res["ocr"] else "없음")
